import json
import logging
import os
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from pathlib import Path

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, 
    PageBreak, Image, KeepTogether, PageTemplate, Frame
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY, TA_RIGHT
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .config import AgentConfig
from .ollama_client import OllamaClient
from .orchestrator import RunResult, TagAnalysisResult
from .solution_verifier import Evidence, VerifiedSuggestion
from .temporal_comparator import ComparisonResult
from .cache_manager import CacheManager

logger = logging.getLogger(__name__)

# Color definitions
COLOR_DARK_BLUE = colors.HexColor("#1A237E")
COLOR_LIGHT_BLUE = colors.HexColor("#E3F2FD")
COLOR_VERIFIED_GREEN = colors.HexColor("#2E7D32")
COLOR_UNVERIFIED_RED = colors.HexColor("#C62828")
COLOR_LIGHT_GRAY = colors.HexColor("#E0E0E0")
COLOR_WHITE = colors.HexColor("#FFFFFF")

# RGB Colors for DOCX
RGB_DARK_BLUE = RGBColor(26, 35, 126)
RGB_VERIFIED_GREEN = RGBColor(46, 125, 50)
RGB_UNVERIFIED_RED = RGBColor(198, 40, 40)


class ReportGenerator:
    """Generates PDF and DOCX reports from StackOverflow analysis results."""

    def __init__(self, config: AgentConfig, llm: OllamaClient, cache: Optional[CacheManager] = None):
        """
        Initialize report generator.

        Args:
            config: AgentConfig instance
            llm: OllamaClient instance
            cache: Optional CacheManager for API log access
        """
        self.config = config
        self.llm = llm
        self.cache = cache

    def generate(self, run_result: RunResult, output_dir: str = "./reports") -> Dict[str, str]:
        """
        Generate both PDF and DOCX reports.

        Args:
            run_result: RunResult object from orchestrator
            output_dir: Output directory for reports

        Returns:
            Dictionary with {"pdf": path, "docx": path}
            If PDF fails, returns only DOCX. If DOCX fails, returns only PDF.
            If both fail, exports JSON fallback.
        """
        # Create output directory
        Path(output_dir).mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_name = f"report_{run_result.run_id[:8]}_{timestamp}"
        
        result = {}
        pdf_path = os.path.join(output_dir, f"{base_name}.pdf")
        docx_path = os.path.join(output_dir, f"{base_name}.docx")

        # Try PDF generation
        pdf_failed = False
        try:
            logger.info(f"Generating PDF report: {pdf_path}")
            self._generate_pdf(run_result, pdf_path)
            result["pdf"] = pdf_path
            logger.info(f"✓ PDF generated successfully")
        except Exception as e:
            pdf_failed = True
            logger.error(f"✗ PDF generation failed: {str(e)}", exc_info=True)

        # Try DOCX generation
        docx_failed = False
        try:
            logger.info(f"Generating DOCX report: {docx_path}")
            self._generate_docx(run_result, docx_path)
            result["docx"] = docx_path
            logger.info(f"✓ DOCX generated successfully")
        except Exception as e:
            docx_failed = True
            logger.error(f"✗ DOCX generation failed: {str(e)}", exc_info=True)

        # Fallback: JSON export
        if pdf_failed and docx_failed:
            json_path = os.path.join(output_dir, f"{base_name}.json")
            try:
                logger.warning("Both PDF and DOCX generation failed. Exporting JSON fallback.")
                self._export_json_fallback(run_result, json_path)
                result["json"] = json_path
            except Exception as e:
                logger.error(f"JSON fallback also failed: {str(e)}", exc_info=True)

        return result

    def _generate_pdf(self, run_result: RunResult, output_path: str) -> None:
        """Generate PDF report using reportlab."""
        from reportlab.platypus import BaseDocTemplate, PageTemplate as PT, Frame
        
        # Create document
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=0.5*inch,
            leftMargin=0.5*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch,
        )

        # Build story
        story = []

        # Cover page
        story.extend(self._cover_page(run_result))
        story.append(PageBreak())

        # Executive summary
        exec_summary = self._generate_executive_summary(run_result)
        story.extend(self._executive_summary(exec_summary, run_result))
        story.append(PageBreak())

        # Methodology
        story.extend(self._methodology(run_result))
        story.append(PageBreak())

        # Findings per tag
        for tag in run_result.tags_analyzed:
            if tag in run_result.tag_analyses:
                tag_result = run_result.tag_analyses[tag]
                story.extend(self._findings_per_tag(tag, tag_result, run_result))
                story.append(PageBreak())

        # Before/after section
        if run_result.comparisons:
            story.extend(self._before_after_section(run_result.comparisons))
            story.append(PageBreak())

        # Knowledge gaps
        unanswered = self._collect_unanswered_questions(run_result)
        if unanswered:
            story.extend(self._knowledge_gap_section(unanswered))
            story.append(PageBreak())

        # Emerging issues
        emerging = self._collect_emerging_issues(run_result)
        if emerging:
            story.extend(self._emerging_issues_section(emerging))
            story.append(PageBreak())

        # Appendices
        all_evidence = self._collect_all_evidence(run_result)
        story.extend(self._appendix_post_ids(all_evidence))
        story.append(PageBreak())

        story.extend(self._appendix_metrics_table(run_result))
        story.append(PageBreak())

        if self.cache:
            story.extend(self._appendix_api_log())

        # Build PDF
        doc.build(story)

    def _generate_docx(self, run_result: RunResult, output_path: str) -> None:
        """Generate DOCX report using python-docx."""
        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Cover page
        self._cover_page_docx(doc, run_result)
        doc.add_page_break()

        # Executive summary
        exec_summary = self._generate_executive_summary(run_result)
        self._executive_summary_docx(doc, exec_summary, run_result)
        doc.add_page_break()

        # Methodology
        self._methodology_docx(doc, run_result)
        doc.add_page_break()

        # Findings per tag
        for tag in run_result.tags_analyzed:
            if tag in run_result.tag_analyses:
                tag_result = run_result.tag_analyses[tag]
                self._findings_per_tag_docx(doc, tag, tag_result, run_result)
                doc.add_page_break()

        # Before/after section
        if run_result.comparisons:
            self._before_after_section_docx(doc, run_result.comparisons)
            doc.add_page_break()

        # Knowledge gaps
        unanswered = self._collect_unanswered_questions(run_result)
        if unanswered:
            self._knowledge_gap_section_docx(doc, unanswered)
            doc.add_page_break()

        # Emerging issues
        emerging = self._collect_emerging_issues(run_result)
        if emerging:
            self._emerging_issues_section_docx(doc, emerging)
            doc.add_page_break()

        # Appendices
        all_evidence = self._collect_all_evidence(run_result)
        self._appendix_post_ids_docx(doc, all_evidence)
        doc.add_page_break()

        self._appendix_metrics_table_docx(doc, run_result)
        doc.add_page_break()

        if self.cache:
            self._appendix_api_log_docx(doc)

        # Save DOCX
        doc.save(output_path)

    def _generate_executive_summary(self, run_result: RunResult) -> str:
        """
        Generate executive summary using LLM.

        If LLM invents metrics not in data, fallback to template.
        """
        key_metrics = self._extract_key_metrics(run_result)
        metrics_json = json.dumps(key_metrics, indent=2)

        prompt = f"""Summarize these findings in 3-5 sentences. Only reference the data provided. Do not invent statistics.

Key Metrics:
{metrics_json}

Write a concise executive summary."""

        system_msg = "You are a technical writer. Write a 3-5 sentence executive summary. Only reference the data provided. Do not invent statistics."

        try:
            summary = self.llm.generate(prompt, system=system_msg, max_tokens=300)
            # Verify summary doesn't contain invented metrics
            if self._verify_summary_accuracy(summary, key_metrics):
                return summary
            else:
                logger.warning("LLM summary contains invented metrics. Using template.")
                return self._template_executive_summary(key_metrics)
        except Exception as e:
            logger.error(f"LLM generation failed: {e}. Using template.")
            return self._template_executive_summary(key_metrics)

    def _verify_summary_accuracy(self, summary: str, key_metrics: Dict) -> bool:
        """Verify that summary only references metrics in key_metrics."""
        # Simple heuristic: check if any numbers in summary exist in metrics
        import re
        numbers = re.findall(r'\d+\.?\d*', summary)
        metric_numbers = set()
        for v in key_metrics.values():
            if isinstance(v, (int, float)):
                metric_numbers.add(str(int(v)))
        
        # If summary has numbers not in metrics, it likely invented them
        for num in numbers:
            if num not in metric_numbers and float(num) not in metric_numbers:
                return False
        return True

    def _template_executive_summary(self, key_metrics: Dict) -> str:
        """Generate template-based executive summary."""
        tags = key_metrics.get("tags_analyzed", 0)
        total_qs = key_metrics.get("total_questions", 0)
        unanswered_pct = key_metrics.get("unanswered_percentage", 0)
        
        return f"""This analysis examined {tags} Stack Overflow tags, covering {total_qs} questions. 
We identified key error patterns and solution clusters across the dataset. 
Approximately {unanswered_pct:.1f}% of questions remain unanswered for extended periods. 
See findings for detailed recommendations and temporal trends."""

    def _extract_key_metrics(self, run_result: RunResult) -> Dict:
        """Extract key metrics from run_result."""
        total_questions = 0
        total_answered = 0
        total_unanswered = 0
        total_clusters = 0

        for tag_result in run_result.tag_analyses.values():
            if tag_result.analysis:
                total_questions += tag_result.analysis.total_questions
                total_answered += tag_result.analysis.answered_count
                total_unanswered += tag_result.analysis.unanswered_count
                total_clusters += len(tag_result.analysis.clusters)

        unanswered_pct = (total_unanswered / total_questions * 100) if total_questions > 0 else 0

        return {
            "tags_analyzed": len(run_result.tags_analyzed),
            "total_questions": total_questions,
            "answered_count": total_answered,
            "unanswered_count": total_unanswered,
            "unanswered_percentage": unanswered_pct,
            "total_clusters": total_clusters,
            "suggestions_count": sum(len(suggs) for suggs in run_result.suggestions.values()),
        }

    # ===== PDF SECTION METHODS =====

    def _cover_page(self, run_result: RunResult) -> List:
        """Generate PDF cover page."""
        styles = getSampleStyleSheet()
        story = []

        # Spacer
        story.append(Spacer(1, 2*inch))

        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=28,
            textColor=COLOR_DARK_BLUE,
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        story.append(Paragraph("Stack Overflow Intelligence Report", title_style))

        # Run ID
        story.append(Spacer(1, 0.3*inch))
        subtitle_style = ParagraphStyle(
            'Subtitle',
            parent=styles['Normal'],
            fontSize=14,
            alignment=TA_CENTER,
            textColor=COLOR_DARK_BLUE,
        )
        story.append(Paragraph(f"Run ID: {run_result.run_id}", subtitle_style))

        # Tags
        story.append(Spacer(1, 0.2*inch))
        tags_str = ", ".join(run_result.tags_analyzed)
        story.append(Paragraph(f"<b>Tags Analyzed:</b> {tags_str}", subtitle_style))

        # Date range
        story.append(Spacer(1, 0.2*inch))
        date_range = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        story.append(Paragraph(date_range, subtitle_style))

        # Status
        story.append(Spacer(1, 0.3*inch))
        status_color = COLOR_VERIFIED_GREEN if run_result.status == "SUCCESS" else COLOR_UNVERIFIED_RED
        status_style = ParagraphStyle(
            'Status',
            parent=styles['Normal'],
            fontSize=16,
            alignment=TA_CENTER,
            textColor=status_color,
            fontName='Helvetica-Bold'
        )
        story.append(Paragraph(f"Status: {run_result.status}", status_style))

        return story

    def _executive_summary(self, exec_summary: str, run_result: RunResult) -> List:
        """PDF executive summary section."""
        styles = getSampleStyleSheet()
        story = []

        # Heading
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=COLOR_DARK_BLUE,
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        story.append(Paragraph("Executive Summary", heading_style))
        story.append(Spacer(1, 0.1*inch))

        # Summary text
        body_style = ParagraphStyle(
            'Body',
            parent=styles['Normal'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=12,
        )
        story.append(Paragraph(exec_summary, body_style))

        # Key metrics
        story.append(Spacer(1, 0.2*inch))
        metrics = self._extract_key_metrics(run_result)
        metrics_data = [
            ["Metric", "Value"],
            ["Tags Analyzed", str(metrics["tags_analyzed"])],
            ["Total Questions", str(metrics["total_questions"])],
            ["Answered", str(metrics["answered_count"])],
            ["Unanswered", str(metrics["unanswered_count"])],
            ["Unanswered %", f"{metrics['unanswered_percentage']:.1f}%"],
            ["Error Clusters", str(metrics["total_clusters"])],
            ["Verified Suggestions", str(metrics["suggestions_count"])],
        ]

        metrics_table = Table(metrics_data, colWidths=[3*inch, 2*inch])
        metrics_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), COLOR_DARK_BLUE),
            ('TEXTCOLOR', (0, 0), (-1, 0), COLOR_WHITE),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), COLOR_LIGHT_GRAY),
            ('GRID', (0, 0), (-1, -1), 1, COLOR_LIGHT_GRAY),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [COLOR_WHITE, COLOR_LIGHT_GRAY]),
        ]))
        story.append(metrics_table)

        return story

    def _methodology(self, run_result: RunResult) -> List:
        """PDF methodology section (no LLM)."""
        styles = getSampleStyleSheet()
        story = []

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=COLOR_DARK_BLUE,
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        story.append(Paragraph("Methodology", heading_style))
        story.append(Spacer(1, 0.1*inch))

        body_style = ParagraphStyle(
            'Body',
            parent=styles['Normal'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=12,
        )

        story.append(Paragraph(
            "This analysis employed the following methodology:",
            body_style
        ))
        story.append(Spacer(1, 0.1*inch))

        # Steps
        steps = [
            "<b>1. Data Collection:</b> Fetched recent questions and answers from Stack Overflow using official API.",
            "<b>2. Pattern Analysis:</b> Clustered questions by error messages and common failure modes.",
            "<b>3. Solution Verification:</b> Ranked answer clusters by acceptance rate, score, and consistency.",
            "<b>4. Confidence Scoring:</b> Evaluated confidence based on answer quality, sample size, and temporal stability.",
            "<b>5. Temporal Comparison:</b> Compared pre/post intervention metrics to measure impact.",
            "<b>6. Report Generation:</b> Synthesized findings with citations to original Stack Overflow posts.",
        ]

        for step in steps:
            story.append(Paragraph(step, body_style))
            story.append(Spacer(1, 0.05*inch))

        story.append(Spacer(1, 0.1*inch))
        story.append(Paragraph(
            f"<b>Run Duration:</b> {run_result.duration_seconds:.1f} seconds | "
            f"<b>Status:</b> {run_result.status}",
            body_style
        ))

        return story

    def _findings_per_tag(self, tag: str, tag_result: TagAnalysisResult, run_result: RunResult) -> List:
        """PDF findings for a specific tag."""
        styles = getSampleStyleSheet()
        story = []

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=COLOR_DARK_BLUE,
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        story.append(Paragraph(f"Findings: {tag.upper()}", heading_style))
        story.append(Spacer(1, 0.1*inch))

        if tag_result.status == "FAILED":
            error_style = ParagraphStyle(
                'Error',
                parent=styles['Normal'],
                fontSize=11,
                textColor=COLOR_UNVERIFIED_RED,
            )
            story.append(Paragraph(f"Analysis failed: {tag_result.error}", error_style))
            return story

        analysis = tag_result.analysis
        body_style = ParagraphStyle(
            'Body',
            parent=styles['Normal'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=12,
        )

        # Summary stats
        if analysis:
            summary = f"""<b>Overview:</b> Analyzed {analysis.total_questions} questions, 
{analysis.answered_count} answered ({100*analysis.answered_count/analysis.total_questions:.1f}%), 
{analysis.unanswered_count} unanswered. Trend: {analysis.trend_direction}. 
{len(analysis.clusters)} error clusters identified."""
            story.append(Paragraph(summary, body_style))
            story.append(Spacer(1, 0.15*inch))

            # Clusters table
            cluster_data = [["Cluster", "Issues", "Trend", "Emerging", "Verified"]]
            for cluster in analysis.clusters:
                verified_icon = "✓" if any(
                    s.cluster_id == cluster.cluster_id for s in tag_result.suggestions
                ) else "✗"
                cluster_data.append([
                    cluster.label[:20],
                    str(cluster.question_count),
                    cluster.trend_direction,
                    "Yes" if cluster.is_emerging else "No",
                    verified_icon
                ])

            if len(cluster_data) > 1:
                cluster_table = Table(cluster_data, colWidths=[2*inch, 1*inch, 1*inch, 0.8*inch, 0.8*inch])
                cluster_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), COLOR_DARK_BLUE),
                    ('TEXTCOLOR', (0, 0), (-1, 0), COLOR_WHITE),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                    ('GRID', (0, 0), (-1, -1), 1, COLOR_LIGHT_GRAY),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [COLOR_WHITE, COLOR_LIGHT_GRAY]),
                ]))
                story.append(cluster_table)
                story.append(Spacer(1, 0.15*inch))

        # Suggestions
        if tag_result.suggestions:
            sugg_heading = ParagraphStyle(
                'SubHeading',
                parent=styles['Heading2'],
                fontSize=12,
                textColor=COLOR_DARK_BLUE,
                spaceAfter=10,
                fontName='Helvetica-Bold'
            )
            story.append(Paragraph("Verified Suggestions", sugg_heading))
            story.append(Spacer(1, 0.1*inch))

            for sugg in tag_result.suggestions:
                status_icon = "✓ VERIFIED" if sugg.status == "VERIFIED" else "⚠ " + sugg.status
                status_color = COLOR_VERIFIED_GREEN if sugg.status == "VERIFIED" else COLOR_UNVERIFIED_RED

                sugg_style = ParagraphStyle(
                    'Suggestion',
                    parent=styles['Normal'],
                    fontSize=10,
                    textColor=status_color,
                    spaceAfter=8,
                )
                story.append(Paragraph(f"<b>{sugg.title}</b> [{status_icon}]", sugg_style))

                summary_style = ParagraphStyle(
                    'SummaryText',
                    parent=styles['Normal'],
                    fontSize=10,
                    alignment=TA_JUSTIFY,
                    spaceAfter=6,
                    leftIndent=0.3*inch,
                )
                story.append(Paragraph(sugg.summary, summary_style))

                # Evidence citations
                evidence_str = ", ".join([f"[Q{e.question_id}]" for e in sugg.evidence])
                story.append(Paragraph(f"<i>Evidence: {evidence_str}</i>", summary_style))
                story.append(Spacer(1, 0.1*inch))

        return story

    def _before_after_section(self, comparisons: Dict[str, ComparisonResult]) -> List:
        """PDF before/after comparison section."""
        styles = getSampleStyleSheet()
        story = []

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=COLOR_DARK_BLUE,
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        story.append(Paragraph("Temporal Impact Analysis", heading_style))
        story.append(Spacer(1, 0.1*inch))

        body_style = ParagraphStyle(
            'Body',
            parent=styles['Normal'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=12,
        )

        for tag, comparison in comparisons.items():
            story.append(Paragraph(f"<b>{tag}</b> (Intervention: {comparison.intervention_date})", body_style))
            story.append(Spacer(1, 0.05*inch))

            # Metrics table
            metrics_data = [
                ["Metric", "Pre", "Post", "Delta", "Trend"],
            ]
            for metric_name, delta in comparison.deltas.items():
                metrics_data.append([
                    metric_name,
                    f"{delta.pre_value:.2f}",
                    f"{delta.post_value:.2f}",
                    f"{delta.delta:+.2f}",
                    delta.trend,
                ])

            metrics_table = Table(metrics_data, colWidths=[2*inch, 1*inch, 1*inch, 1*inch, 0.5*inch])
            metrics_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), COLOR_DARK_BLUE),
                ('TEXTCOLOR', (0, 0), (-1, 0), COLOR_WHITE),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('GRID', (0, 0), (-1, -1), 1, COLOR_LIGHT_GRAY),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [COLOR_WHITE, COLOR_LIGHT_GRAY]),
            ]))
            story.append(metrics_table)

            # Verdict
            verdict_color = COLOR_VERIFIED_GREEN if comparison.verdict == "IMPROVED" else COLOR_UNVERIFIED_RED
            verdict_style = ParagraphStyle(
                'Verdict',
                parent=styles['Normal'],
                fontSize=10,
                textColor=verdict_color,
                spaceAfter=12,
            )
            story.append(Spacer(1, 0.1*inch))
            story.append(Paragraph(f"<b>Verdict:</b> {comparison.verdict}", verdict_style))
            story.append(Spacer(1, 0.15*inch))

        return story

    def _knowledge_gap_section(self, unanswered: List[Dict]) -> List:
        """PDF knowledge gap section."""
        styles = getSampleStyleSheet()
        story = []

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=COLOR_DARK_BLUE,
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        story.append(Paragraph("Knowledge Gaps (7+ days unanswered)", heading_style))
        story.append(Spacer(1, 0.1*inch))

        body_style = ParagraphStyle(
            'Body',
            parent=styles['Normal'],
            fontSize=10,
            alignment=TA_JUSTIFY,
            spaceAfter=8,
        )

        story.append(Paragraph(
            f"Identified {len(unanswered)} questions unanswered for 7+ days. These represent potential gaps in community knowledge.",
            body_style
        ))
        story.append(Spacer(1, 0.1*inch))

        for item in unanswered[:10]:  # Limit to first 10
            title = item.get("title", "Unknown")
            url = item.get("url", "#")
            tag = item.get("tag", "N/A")
            story.append(Paragraph(f"<b>Q{item.get('question_id', 0)}:</b> {title} [tag: {tag}]", body_style))

        if len(unanswered) > 10:
            story.append(Paragraph(f"... and {len(unanswered) - 10} more", body_style))

        return story

    def _emerging_issues_section(self, emerging_clusters: List) -> List:
        """PDF emerging issues section."""
        styles = getSampleStyleSheet()
        story = []

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=COLOR_DARK_BLUE,
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        story.append(Paragraph("Emerging Issues", heading_style))
        story.append(Spacer(1, 0.1*inch))

        body_style = ParagraphStyle(
            'Body',
            parent=styles['Normal'],
            fontSize=10,
            alignment=TA_JUSTIFY,
            spaceAfter=8,
        )

        story.append(Paragraph(
            f"Detected {len(emerging_clusters)} emerging error patterns showing increasing trend.",
            body_style
        ))
        story.append(Spacer(1, 0.1*inch))

        for cluster in emerging_clusters:
            story.append(Paragraph(
                f"<b>{cluster.get('label', 'Unknown')}</b>: {cluster.get('question_count', 0)} issues, "
                f"trend {cluster.get('trend_direction', 'UNKNOWN')}",
                body_style
            ))

        return story

    def _appendix_post_ids(self, all_evidence: List[Evidence]) -> List:
        """PDF appendix with all referenced post IDs."""
        styles = getSampleStyleSheet()
        story = []

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=14,
            textColor=COLOR_DARK_BLUE,
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        story.append(Paragraph("APPENDIX A: Source Documentation", heading_style))
        story.append(Spacer(1, 0.1*inch))

        body_style = ParagraphStyle(
            'Body',
            parent=styles['Normal'],
            fontSize=9,
            spaceAfter=6,
        )

        # Evidence table
        evidence_data = [["Question ID", "Answer ID", "Score", "Accepted", "URL"]]
        for evidence in all_evidence:
            evidence_data.append([
                str(evidence.question_id),
                str(evidence.answer_id),
                str(evidence.answer_score),
                "Yes" if evidence.is_accepted else "No",
                evidence.question_url[:30] + "..." if len(evidence.question_url) > 30 else evidence.question_url,
            ])

        if len(evidence_data) > 1:
            evidence_table = Table(evidence_data, colWidths=[1*inch, 1*inch, 0.7*inch, 0.8*inch, 1.5*inch])
            evidence_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), COLOR_DARK_BLUE),
                ('TEXTCOLOR', (0, 0), (-1, 0), COLOR_WHITE),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, COLOR_LIGHT_GRAY),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [COLOR_WHITE, COLOR_LIGHT_GRAY]),
            ]))
            story.append(evidence_table)

        return story

    def _appendix_metrics_table(self, run_result: RunResult) -> List:
        """PDF appendix with metrics table."""
        styles = getSampleStyleSheet()
        story = []

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=14,
            textColor=COLOR_DARK_BLUE,
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        story.append(Paragraph("APPENDIX B: Detailed Metrics", heading_style))
        story.append(Spacer(1, 0.1*inch))

        # Metrics table
        metrics_data = [["Tag", "Questions", "Answered", "Unanswered", "Clusters", "Trend"]]
        for tag in run_result.tags_analyzed:
            if tag in run_result.tag_analyses:
                tag_result = run_result.tag_analyses[tag]
                if tag_result.analysis:
                    analysis = tag_result.analysis
                    metrics_data.append([
                        tag,
                        str(analysis.total_questions),
                        str(analysis.answered_count),
                        str(analysis.unanswered_count),
                        str(len(analysis.clusters)),
                        analysis.trend_direction,
                    ])

        if len(metrics_data) > 1:
            metrics_table = Table(metrics_data, colWidths=[1.5*inch, 1*inch, 1*inch, 1*inch, 1*inch, 1*inch])
            metrics_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), COLOR_DARK_BLUE),
                ('TEXTCOLOR', (0, 0), (-1, 0), COLOR_WHITE),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 1, COLOR_LIGHT_GRAY),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [COLOR_WHITE, COLOR_LIGHT_GRAY]),
            ]))
            story.append(metrics_table)

        return story

    def _appendix_api_log(self) -> List:
        """PDF appendix with API call log."""
        styles = getSampleStyleSheet()
        story = []

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=14,
            textColor=COLOR_DARK_BLUE,
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        story.append(Paragraph("APPENDIX C: API Call Log", heading_style))
        story.append(Spacer(1, 0.1*inch))

        if not self.cache:
            body_style = ParagraphStyle(
                'Body',
                parent=styles['Normal'],
                fontSize=10,
            )
            story.append(Paragraph("No cache available.", body_style))
            return story

        # Get API logs
        try:
            logs = self.cache.db["api_call_log"].rows_where(order_by="-called_at", limit=20)
            log_data = [["Endpoint", "Tag", "Page", "Quota", "Duration (ms)"]]
            for log in logs:
                log_data.append([
                    log["endpoint"],
                    log["tag"],
                    str(log["page"]),
                    str(log["quota_remaining"]),
                    str(log["duration_ms"]),
                ])

            if len(log_data) > 1:
                log_table = Table(log_data, colWidths=[1.5*inch, 1.5*inch, 0.7*inch, 1*inch, 1*inch])
                log_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), COLOR_DARK_BLUE),
                    ('TEXTCOLOR', (0, 0), (-1, 0), COLOR_WHITE),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 0.5, COLOR_LIGHT_GRAY),
                ]))
                story.append(log_table)
        except Exception as e:
            logger.error(f"Error retrieving API logs: {e}")

        return story

    # ===== DOCX SECTION METHODS =====

    def _cover_page_docx(self, doc: Document, run_result: RunResult) -> None:
        """DOCX cover page."""
        # Title
        title = doc.add_paragraph()
        title_run = title.add_run("Stack Overflow Intelligence Report")
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGB_DARK_BLUE
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(30)

        # Run ID
        doc.add_paragraph()
        run_id_p = doc.add_paragraph(f"Run ID: {run_result.run_id}")
        run_id_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Tags
        tags_p = doc.add_paragraph(f"Tags Analyzed: {', '.join(run_result.tags_analyzed)}")
        tags_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Date
        date_p = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Status
        doc.add_paragraph()
        status_p = doc.add_paragraph(f"Status: {run_result.status}")
        status_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        status_run = status_p.runs[0]
        status_run.font.bold = True
        status_run.font.size = Pt(14)
        status_run.font.color.rgb = RGB_VERIFIED_GREEN if run_result.status == "SUCCESS" else RGB_UNVERIFIED_RED

    def _executive_summary_docx(self, doc: Document, exec_summary: str, run_result: RunResult) -> None:
        """DOCX executive summary."""
        heading = doc.add_heading("Executive Summary", level=1)
        heading.runs[0].font.color.rgb = RGB_DARK_BLUE

        doc.add_paragraph(exec_summary)

        # Metrics table
        metrics = self._extract_key_metrics(run_result)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Metric"
        hdr_cells[1].text = "Value"
        
        rows_data = [
            ["Tags Analyzed", str(metrics["tags_analyzed"])],
            ["Total Questions", str(metrics["total_questions"])],
            ["Answered", str(metrics["answered_count"])],
            ["Unanswered", str(metrics["unanswered_count"])],
            ["Unanswered %", f"{metrics['unanswered_percentage']:.1f}%"],
            ["Error Clusters", str(metrics["total_clusters"])],
            ["Verified Suggestions", str(metrics["suggestions_count"])],
        ]

        for row_data in rows_data:
            row_cells = table.add_row().cells
            row_cells[0].text = row_data[0]
            row_cells[1].text = row_data[1]

    def _methodology_docx(self, doc: Document, run_result: RunResult) -> None:
        """DOCX methodology section."""
        heading = doc.add_heading("Methodology", level=1)
        heading.runs[0].font.color.rgb = RGB_DARK_BLUE

        doc.add_paragraph("This analysis employed the following methodology:")

        steps = [
            "1. Data Collection: Fetched recent questions and answers from Stack Overflow using official API.",
            "2. Pattern Analysis: Clustered questions by error messages and common failure modes.",
            "3. Solution Verification: Ranked answer clusters by acceptance rate, score, and consistency.",
            "4. Confidence Scoring: Evaluated confidence based on answer quality, sample size, and temporal stability.",
            "5. Temporal Comparison: Compared pre/post intervention metrics to measure impact.",
            "6. Report Generation: Synthesized findings with citations to original Stack Overflow posts.",
        ]

        for step in steps:
            doc.add_paragraph(step)

        doc.add_paragraph()
        doc.add_paragraph(
            f"Run Duration: {run_result.duration_seconds:.1f} seconds | Status: {run_result.status}"
        )

    def _findings_per_tag_docx(self, doc: Document, tag: str, tag_result: TagAnalysisResult, run_result: RunResult) -> None:
        """DOCX findings for a specific tag."""
        heading = doc.add_heading(f"Findings: {tag.upper()}", level=1)
        heading.runs[0].font.color.rgb = RGB_DARK_BLUE

        if tag_result.status == "FAILED":
            doc.add_paragraph(f"Analysis failed: {tag_result.error}")
            return

        analysis = tag_result.analysis
        if analysis:
            summary = f"""Overview: Analyzed {analysis.total_questions} questions, 
{analysis.answered_count} answered ({100*analysis.answered_count/analysis.total_questions:.1f}%), 
{analysis.unanswered_count} unanswered. Trend: {analysis.trend_direction}. 
{len(analysis.clusters)} error clusters identified."""
            doc.add_paragraph(summary)

            # Clusters table
            if analysis.clusters:
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Light Grid Accent 1'
                
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Cluster"
                hdr_cells[1].text = "Issues"
                hdr_cells[2].text = "Trend"
                hdr_cells[3].text = "Emerging"
                hdr_cells[4].text = "Verified"

                for cluster in analysis.clusters:
                    verified_icon = "✓" if any(
                        s.cluster_id == cluster.cluster_id for s in tag_result.suggestions
                    ) else "✗"
                    row_cells = table.add_row().cells
                    row_cells[0].text = cluster.label[:20]
                    row_cells[1].text = str(cluster.question_count)
                    row_cells[2].text = cluster.trend_direction
                    row_cells[3].text = "Yes" if cluster.is_emerging else "No"
                    row_cells[4].text = verified_icon

        # Suggestions
        if tag_result.suggestions:
            doc.add_heading("Verified Suggestions", level=2)
            for sugg in tag_result.suggestions:
                status_icon = "✓ VERIFIED" if sugg.status == "VERIFIED" else "⚠ " + sugg.status
                p = doc.add_paragraph(f"{sugg.title} [{status_icon}]")
                p_format = p.paragraph_format
                p_format.left_indent = Inches(0.25)

                doc.add_paragraph(sugg.summary, style='List Bullet')
                evidence_str = ", ".join([f"[Q{e.question_id}]" for e in sugg.evidence])
                doc.add_paragraph(f"Evidence: {evidence_str}", style='List Bullet')

    def _before_after_section_docx(self, doc: Document, comparisons: Dict[str, ComparisonResult]) -> None:
        """DOCX before/after comparison."""
        heading = doc.add_heading("Temporal Impact Analysis", level=1)
        heading.runs[0].font.color.rgb = RGB_DARK_BLUE

        for tag, comparison in comparisons.items():
            doc.add_heading(f"{tag} (Intervention: {comparison.intervention_date})", level=2)

            # Metrics table
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Light Grid Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Metric"
            hdr_cells[1].text = "Pre"
            hdr_cells[2].text = "Post"
            hdr_cells[3].text = "Delta"
            hdr_cells[4].text = "Trend"

            for metric_name, delta in comparison.deltas.items():
                row_cells = table.add_row().cells
                row_cells[0].text = metric_name
                row_cells[1].text = f"{delta.pre_value:.2f}"
                row_cells[2].text = f"{delta.post_value:.2f}"
                row_cells[3].text = f"{delta.delta:+.2f}"
                row_cells[4].text = delta.trend

            # Verdict
            doc.add_paragraph(f"Verdict: {comparison.verdict}")

    def _knowledge_gap_section_docx(self, doc: Document, unanswered: List[Dict]) -> None:
        """DOCX knowledge gap section."""
        heading = doc.add_heading("Knowledge Gaps (7+ days unanswered)", level=1)
        heading.runs[0].font.color.rgb = RGB_DARK_BLUE

        doc.add_paragraph(
            f"Identified {len(unanswered)} questions unanswered for 7+ days. These represent potential gaps in community knowledge."
        )

        for item in unanswered[:10]:
            title = item.get("title", "Unknown")
            tag = item.get("tag", "N/A")
            doc.add_paragraph(
                f"Q{item.get('question_id', 0)}: {title} [tag: {tag}]",
                style='List Bullet'
            )

        if len(unanswered) > 10:
            doc.add_paragraph(f"... and {len(unanswered) - 10} more")

    def _emerging_issues_section_docx(self, doc: Document, emerging_clusters: List) -> None:
        """DOCX emerging issues section."""
        heading = doc.add_heading("Emerging Issues", level=1)
        heading.runs[0].font.color.rgb = RGB_DARK_BLUE

        doc.add_paragraph(
            f"Detected {len(emerging_clusters)} emerging error patterns showing increasing trend."
        )

        for cluster in emerging_clusters:
            doc.add_paragraph(
                f"{cluster.get('label', 'Unknown')}: {cluster.get('question_count', 0)} issues, "
                f"trend {cluster.get('trend_direction', 'UNKNOWN')}",
                style='List Bullet'
            )

    def _appendix_post_ids_docx(self, doc: Document, all_evidence: List[Evidence]) -> None:
        """DOCX appendix with post IDs."""
        heading = doc.add_heading("APPENDIX A: Source Documentation", level=1)
        heading.runs[0].font.color.rgb = RGB_DARK_BLUE

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Question ID"
        hdr_cells[1].text = "Answer ID"
        hdr_cells[2].text = "Score"
        hdr_cells[3].text = "Accepted"
        hdr_cells[4].text = "URL"

        for evidence in all_evidence:
            row_cells = table.add_row().cells
            row_cells[0].text = str(evidence.question_id)
            row_cells[1].text = str(evidence.answer_id)
            row_cells[2].text = str(evidence.answer_score)
            row_cells[3].text = "Yes" if evidence.is_accepted else "No"
            row_cells[4].text = evidence.question_url[:50] + "..." if len(evidence.question_url) > 50 else evidence.question_url

    def _appendix_metrics_table_docx(self, doc: Document, run_result: RunResult) -> None:
        """DOCX appendix with metrics."""
        heading = doc.add_heading("APPENDIX B: Detailed Metrics", level=1)
        heading.runs[0].font.color.rgb = RGB_DARK_BLUE

        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Tag"
        hdr_cells[1].text = "Questions"
        hdr_cells[2].text = "Answered"
        hdr_cells[3].text = "Unanswered"
        hdr_cells[4].text = "Clusters"
        hdr_cells[5].text = "Trend"

        for tag in run_result.tags_analyzed:
            if tag in run_result.tag_analyses:
                tag_result = run_result.tag_analyses[tag]
                if tag_result.analysis:
                    analysis = tag_result.analysis
                    row_cells = table.add_row().cells
                    row_cells[0].text = tag
                    row_cells[1].text = str(analysis.total_questions)
                    row_cells[2].text = str(analysis.answered_count)
                    row_cells[3].text = str(analysis.unanswered_count)
                    row_cells[4].text = str(len(analysis.clusters))
                    row_cells[5].text = analysis.trend_direction

    def _appendix_api_log_docx(self, doc: Document) -> None:
        """DOCX appendix with API log."""
        heading = doc.add_heading("APPENDIX C: API Call Log", level=1)
        heading.runs[0].font.color.rgb = RGB_DARK_BLUE

        if not self.cache:
            doc.add_paragraph("No cache available.")
            return

        try:
            logs = list(self.cache.db["api_call_log"].rows_where(order_by="-called_at", limit=20))
            
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Light Grid Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Endpoint"
            hdr_cells[1].text = "Tag"
            hdr_cells[2].text = "Page"
            hdr_cells[3].text = "Quota"
            hdr_cells[4].text = "Duration (ms)"

            for log in logs:
                row_cells = table.add_row().cells
                row_cells[0].text = log["endpoint"]
                row_cells[1].text = log["tag"]
                row_cells[2].text = str(log["page"])
                row_cells[3].text = str(log["quota_remaining"])
                row_cells[4].text = str(log["duration_ms"])
        except Exception as e:
            logger.error(f"Error retrieving API logs: {e}")
            doc.add_paragraph(f"Error retrieving logs: {str(e)}")

    # ===== HELPER METHODS =====

    def _collect_unanswered_questions(self, run_result: RunResult) -> List[Dict]:
        """Collect all unanswered questions."""
        unanswered = []
        for tag_result in run_result.tag_analyses.values():
            if tag_result.analysis:
                for q in tag_result.analysis.unanswered_7d_plus:
                    unanswered.append({
                        "question_id": q.get("question_id", 0),
                        "title": q.get("title", ""),
                        "url": q.get("link", ""),
                        "tag": tag_result.tag,
                    })
        return unanswered

    def _collect_emerging_issues(self, run_result: RunResult) -> List:
        """Collect all emerging issue clusters."""
        emerging = []
        for tag_result in run_result.tag_analyses.values():
            if tag_result.analysis:
                for cluster in tag_result.analysis.clusters:
                    if cluster.is_emerging:
                        emerging.append({
                            "label": cluster.label,
                            "question_count": cluster.question_count,
                            "trend_direction": cluster.trend_direction,
                            "tag": tag_result.tag,
                        })
        return emerging

    def _collect_all_evidence(self, run_result: RunResult) -> List[Evidence]:
        """Collect all evidence from suggestions."""
        all_evidence = []
        for suggestions in run_result.suggestions.values():
            for sugg in suggestions:
                all_evidence.extend(sugg.evidence)
        return all_evidence

    def _export_json_fallback(self, run_result: RunResult, output_path: str) -> None:
        """Export run_result as JSON fallback."""
        from dataclasses import asdict

        def serialize(obj):
            if hasattr(obj, '__dataclass_fields__'):
                return asdict(obj)
            elif hasattr(obj, '__dict__'):
                return obj.__dict__
            else:
                return str(obj)

        data = {
            "run_id": run_result.run_id,
            "status": run_result.status,
            "tags_analyzed": run_result.tags_analyzed,
            "duration_seconds": run_result.duration_seconds,
            "errors": run_result.errors,
            "warnings": run_result.warnings,
            "tag_analyses": {
                tag: serialize(result) for tag, result in run_result.tag_analyses.items()
            },
            "suggestions": {
                tag: [serialize(s) for s in suggs] 
                for tag, suggs in run_result.suggestions.items()
            },
            "comparisons": {
                tag: serialize(comp) for tag, comp in (run_result.comparisons or {}).items()
            },
        }

        with open(output_path, 'w') as f:
            json.dump(data, f, indent=2, default=str)

        logger.info(f"✓ JSON fallback exported: {output_path}")
